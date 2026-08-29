"""Testes de RAG-023: extração de conteúdo via `DoclingDocumentParser`.

Cobre os critérios de aceite da atividade: extrai texto e metadados
para os formatos suportados; erro de parsing é categorizado
(`UnsupportedDocumentFormatError` vs. `DocumentParsingError`); fixtures
cobrem os quatro formatos aceitos no upload (RAG-021) — PDF incluído,
mesmo sem extração real (ver docstring de `adapters.docling.parser`
sobre por que PDF ainda não é suportado por este adapter).
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from adapters.docling.parser import DoclingDocumentParser
from packages.application.ports.document_parser import (
    DocumentParsingError,
    ParsedDocument,
    UnsupportedDocumentFormatError,
)


def _make_docx_bytes(*, heading: str, paragraph: str) -> bytes:
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def parser() -> DoclingDocumentParser:
    return DoclingDocumentParser()


async def test_parses_markdown(parser: DoclingDocumentParser) -> None:
    content = "# Título\n\nAlgum texto.\n\n## Seção 2\n\nMais texto.".encode()

    result = await parser.parse(filename="doc.md", content=content, content_type="text/markdown")

    assert isinstance(result, ParsedDocument)
    assert "Título" in result.markdown
    assert "Seção 2" in result.markdown
    assert result.original_mimetype == "text/markdown"
    assert result.page_count is None


async def test_parses_plain_text_as_markdown(parser: DoclingDocumentParser) -> None:
    content = "Linha única sem sintaxe especial de markdown.".encode()

    result = await parser.parse(filename="notas.txt", content=content, content_type="text/plain")

    assert "Linha única sem sintaxe especial" in result.markdown
    assert result.original_mimetype == "text/plain"


async def test_parses_docx(parser: DoclingDocumentParser) -> None:
    content = _make_docx_bytes(heading="Título do documento", paragraph="Parágrafo normal.")
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    result = await parser.parse(filename="doc.docx", content=content, content_type=content_type)

    assert "Título do documento" in result.markdown
    assert "Parágrafo normal." in result.markdown
    assert result.original_mimetype == content_type


async def test_pdf_raises_unsupported_format_error(parser: DoclingDocumentParser) -> None:
    # Não precisa ser um PDF válido: o adapter rejeita application/pdf
    # antes de tentar decodificar o conteúdo (ver docstring do módulo
    # sobre por que PDF ainda não é suportado).
    content = b"%PDF-1.4 conteudo nao processado por este adapter"

    with pytest.raises(UnsupportedDocumentFormatError) as exc_info:
        await parser.parse(filename="doc.pdf", content=content, content_type="application/pdf")

    assert exc_info.value.content_type == "application/pdf"


async def test_unknown_content_type_raises_unsupported_format_error(
    parser: DoclingDocumentParser,
) -> None:
    with pytest.raises(UnsupportedDocumentFormatError) as exc_info:
        await parser.parse(
            filename="doc.bin", content=b"x", content_type="application/octet-stream"
        )

    assert exc_info.value.content_type == "application/octet-stream"


async def test_malformed_docx_raises_categorized_parsing_error(
    parser: DoclingDocumentParser,
) -> None:
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    with pytest.raises(DocumentParsingError) as exc_info:
        await parser.parse(
            filename="corrompido.docx",
            content=b"isto nao e um docx valido",
            content_type=content_type,
        )

    assert not isinstance(exc_info.value, UnsupportedDocumentFormatError)
    assert exc_info.value.content_type == content_type
